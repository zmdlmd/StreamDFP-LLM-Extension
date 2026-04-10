#!/usr/bin/env python3
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def _iter_paragraphs(doc: Document):
    for paragraph in doc.paragraphs:
        yield paragraph


def _remove_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def _get_style(doc: Document, name: str):
    for style in doc.styles:
        if style.name == name:
            return style
    return doc.styles[name]


def _clear_paragraph(paragraph) -> None:
    element = paragraph._element
    for child in list(element):
        if child.tag != qn("w:pPr"):
            element.remove(child)


def _add_field(paragraph, instruction: str, placeholder: str = "") -> None:
    run = paragraph.add_run()

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction

    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")

    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_separate)

    if placeholder:
        text = OxmlElement("w:t")
        text.text = placeholder
        run._r.append(text)

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)


def _strip_front_matter(doc: Document) -> None:
    abstract_para = None
    for paragraph in _iter_paragraphs(doc):
        if paragraph.text.strip() == "Abstract":
            abstract_para = paragraph
            break

    if abstract_para is None:
        return

    to_remove = []
    for paragraph in _iter_paragraphs(doc):
        if paragraph._element == abstract_para._element:
            break
        to_remove.append(paragraph)

    for paragraph in to_remove:
        _remove_paragraph(paragraph)


def _insert_toc(doc: Document) -> None:
    anchor = None
    for paragraph in _iter_paragraphs(doc):
        if paragraph.text.strip() == "1. Introduction":
            anchor = paragraph
            break

    if anchor is None:
        return

    toc_heading = anchor.insert_paragraph_before("Table of Contents")
    toc_heading.style = _get_style(doc, "Heading 1")

    toc_para = anchor.insert_paragraph_before("")
    _add_field(
        toc_para,
        'TOC \\o "1-3" \\h \\z \\u',
        "Right-click and choose Update Field to refresh the table of contents.",
    )

    spacer = anchor.insert_paragraph_before("")
    spacer.style = _get_style(doc, "Body Text")


def _style_table_captions(doc: Document) -> None:
    for paragraph in _iter_paragraphs(doc):
        text = paragraph.text.strip()
        if text.startswith("Table ") and "." in text:
            paragraph.style = _get_style(doc, "Table Caption")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_page_numbers(doc: Document) -> None:
    for section in doc.sections:
        footer = section.footer
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _clear_paragraph(paragraph)
        _add_field(paragraph, "PAGE")


def main() -> int:
    if len(sys.argv) < 2:
        raise SystemExit("usage: postprocess_report_docx.py <docx-path>")

    path = Path(sys.argv[1]).resolve()
    doc = Document(path)

    _strip_front_matter(doc)
    _insert_toc(doc)
    _style_table_captions(doc)
    _add_page_numbers(doc)

    doc.save(path)
    print(path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
